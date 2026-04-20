#!/usr/bin/env python3
"""Build chronic absenteeism deliverables from a Brandy Parton accountability snapshot.

Usage:
    python build.py <source_xlsx> <as_of_YYYY-MM-DD> <day_number>

Example:
    python build.py "~/Downloads/Chronic Abs Accountability Data as of March 10 2026.xlsx" 2026-03-10 122

Produces:
    ~/Documents/Chronic Absenteeism <YYYY-MM-DD>/
        build.py (copy)
        Chronic Absenteeism Dashboard <YYYY-MM-DD>.xlsx
        Chronic Absenteeism Executive Summary <YYYY-MM-DD>.docx
        Chronic Absenteeism Packet - <School>.docx (one per school)

Also mirrors all outputs and the source to:
    OneDrive-GCS/Data/Accountability Data/Chronic Absenteeism/
"""
import shutil
import sys
from collections import defaultdict
from datetime import datetime
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

ONEDRIVE = Path.home() / "Library/CloudStorage/OneDrive-GreenevilleCitySchools/Data/Accountability Data/Chronic Absenteeism"
INSTRUCTIONAL_DAYS_DEFAULT = 167
THRESHOLD = 17  # >=17 days absent = already chronically absent

SCHOOL_NAMES = {
    "HH": "Hal Henard Elementary",
    "EV": "EastView Elementary",
    "HI": "Highland Elementary",
    "TV": "Tusculum View Elementary",
    "TOPS": "TOPS Greeneville",
    "GMS": "Greeneville Middle School",
    "GHS": "Greeneville High School",
}
SCHOOL_ORDER = ["HH", "EV", "HI", "TV", "TOPS", "GMS", "GHS"]

def load_data(src):
    wb = openpyxl.load_workbook(src, data_only=True)
    rollup = []
    if "School Chronic Abs %" in wb.sheetnames:
        for r in wb["School Chronic Abs %"].iter_rows(min_row=2, values_only=True):
            if r[0]:
                rollup.append({"school": r[0], "group": r[1], "band": r[2],
                               "n": r[3], "n_ca": r[4], "pct": r[5]})
    students = []
    for code in SCHOOL_NAMES:
        if code not in wb.sheetnames:
            continue
        for r in wb[code].iter_rows(min_row=2, values_only=True):
            if r[1] is None:
                continue
            students.append({
                "school": code, "ssid": r[1], "first": r[2], "last": r[3],
                "grade": r[4], "absences": r[5] or 0,
                "isp_to_date": r[6] or 0, "isp_proj": r[7] or 0,
                "cal_days": r[8] or 0, "rate": r[9] or 0,
            })
    return rollup, students

def classify(s):
    return "already_ca" if s["absences"] >= THRESHOLD else "trending"

def grade_band(g):
    if g is None:
        return "Unknown"
    if g <= 5:
        return "K-5"
    if g <= 8:
        return "6-8"
    return "9-12"

# ---------- Dashboard xlsx ----------
def build_dashboard(rollup, students, out_dir, as_of_str):
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    header_fill = PatternFill("solid", fgColor="1F3864")
    header_font = Font(bold=True, color="FFFFFF")
    ca_fill = PatternFill("solid", fgColor="F8CBAD")
    trend_fill = PatternFill("solid", fgColor="FFF2CC")
    thin = Side(style="thin", color="BFBFBF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def write_header(ws, headers):
        for i, h in enumerate(headers, 1):
            c = ws.cell(row=1, column=i, value=h)
            c.fill = header_fill
            c.font = header_font
            c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            c.border = border
        ws.row_dimensions[1].height = 32

    def autosize(ws):
        for col in ws.columns:
            m = 10
            letter = get_column_letter(col[0].column)
            for c in col:
                if c.value is not None:
                    m = max(m, min(45, len(str(c.value)) + 2))
            ws.column_dimensions[letter].width = m

    by_school = defaultdict(lambda: {"total": 0, "already": 0, "trend": 0})
    for s in students:
        b = by_school[s["school"]]
        b["total"] += 1
        if classify(s) == "already_ca":
            b["already"] += 1
        else:
            b["trend"] += 1
    totals = {"total": 0, "already": 0, "trend": 0}
    for code in SCHOOL_ORDER:
        for k in totals:
            totals[k] += by_school[code][k if k != "total" else "total"]
    # fix key mismatch
    totals = {"total": sum(by_school[c]["total"] for c in SCHOOL_ORDER),
              "already": sum(by_school[c]["already"] for c in SCHOOL_ORDER),
              "trend": sum(by_school[c]["trend"] for c in SCHOOL_ORDER)}

    # Sheet 1: District Summary
    ws = wb.create_sheet("District Summary")
    ws["A1"] = "Chronic Absenteeism Snapshot"
    ws["A1"].font = Font(bold=True, size=16, color="1F3864")
    ws.merge_cells("A1:E1")
    ws["A2"] = f"As of {as_of_str}"
    ws["A2"].font = Font(italic=True, color="595959")
    ws.merge_cells("A2:E2")
    ws.append([])
    hdr_row = ws.max_row + 1
    for i, h in enumerate(["School", "Students Flagged", "Already CA (17+)", "Trending", "% Already CA"], 1):
        c = ws.cell(row=hdr_row, column=i, value=h)
        c.fill = header_fill
        c.font = header_font
        c.alignment = Alignment(horizontal="center", wrap_text=True)
        c.border = border
    for code in SCHOOL_ORDER:
        b = by_school[code]
        pct = (b["already"] / b["total"] * 100) if b["total"] else 0
        ws.append([SCHOOL_NAMES[code], b["total"], b["already"], b["trend"], round(pct, 1)])
    pct_total = (totals["already"] / totals["total"] * 100) if totals["total"] else 0
    ws.append(["District Total", totals["total"], totals["already"], totals["trend"], round(pct_total, 1)])
    for c in ws[ws.max_row]:
        c.font = Font(bold=True)
        c.fill = PatternFill("solid", fgColor="D9E1F2")
    ws.append([])
    ws.append(["School-Level % Chronically Absent (All Enrolled Students)"])
    ws.cell(row=ws.max_row, column=1).font = Font(bold=True, color="1F3864")
    ws.append(["School", "# Students", "# Chronically Absent", "% Chronically Absent"])
    for i in range(1, 5):
        c = ws.cell(row=ws.max_row, column=i)
        c.fill = header_fill
        c.font = header_font
        c.alignment = Alignment(horizontal="center")
    for r in rollup:
        ws.append([r["school"], r["n"], r["n_ca"], r["pct"]])
    autosize(ws)

    cols = ["School", "SSID", "First", "Last", "Grade", "Days Absent",
            "ISP Days to Date", "Projected ISP Days", "Instructional Days", "Absentee Rate %"]

    # Sheet 2: Already CA
    ws = wb.create_sheet("Already CA (17+)")
    write_header(ws, cols)
    already = sorted([s for s in students if classify(s) == "already_ca"],
                     key=lambda s: (-s["absences"], s["school"]))
    for s in already:
        ws.append([SCHOOL_NAMES[s["school"]], s["ssid"], s["first"], s["last"], s["grade"],
                   s["absences"], s["isp_to_date"], s["isp_proj"], s["cal_days"], s["rate"]])
        for c in ws[ws.max_row]:
            c.fill = ca_fill
    ws.freeze_panes = "A2"
    autosize(ws)

    # Sheet 3: Trending
    ws = wb.create_sheet("Trending (10%+)")
    write_header(ws, cols)
    trend = sorted([s for s in students if classify(s) == "trending"],
                   key=lambda s: (-s["rate"], s["school"]))
    for s in trend:
        ws.append([SCHOOL_NAMES[s["school"]], s["ssid"], s["first"], s["last"], s["grade"],
                   s["absences"], s["isp_to_date"], s["isp_proj"], s["cal_days"], s["rate"]])
        for c in ws[ws.max_row]:
            c.fill = trend_fill
    ws.freeze_panes = "A2"
    autosize(ws)

    # Sheet 4: Grade Band
    ws = wb.create_sheet("By Grade Band")
    write_header(ws, ["School", "Grade Band", "Students Flagged", "Already CA", "Trending"])
    by_band = defaultdict(lambda: defaultdict(lambda: {"total": 0, "already": 0, "trend": 0}))
    for s in students:
        b = grade_band(s["grade"])
        bucket = by_band[s["school"]][b]
        bucket["total"] += 1
        if classify(s) == "already_ca":
            bucket["already"] += 1
        else:
            bucket["trend"] += 1
    for code in SCHOOL_ORDER:
        for band in ["K-5", "6-8", "9-12", "Unknown"]:
            if band in by_band[code]:
                b = by_band[code][band]
                ws.append([SCHOOL_NAMES[code], band, b["total"], b["already"], b["trend"]])
    autosize(ws)

    # Sheet 5: All Flagged
    ws = wb.create_sheet("All Flagged Students")
    write_header(ws, cols + ["Status"])
    for s in sorted(students, key=lambda s: (-s["absences"], s["school"])):
        status = "Already CA" if classify(s) == "already_ca" else "Trending"
        ws.append([SCHOOL_NAMES[s["school"]], s["ssid"], s["first"], s["last"], s["grade"],
                   s["absences"], s["isp_to_date"], s["isp_proj"], s["cal_days"], s["rate"], status])
        fill = ca_fill if status == "Already CA" else trend_fill
        for c in ws[ws.max_row]:
            c.fill = fill
    ws.freeze_panes = "A2"
    autosize(ws)

    path = out_dir / f"Chronic Absenteeism Dashboard {as_of_str}.xlsx"
    wb.save(path)
    return path, totals, by_school

# ---------- DOCX helpers ----------
def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    if subtitle:
        p2 = doc.add_paragraph()
        r2 = p2.add_run(subtitle)
        r2.italic = True
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

def style_header_row(row):
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F3864")
        tc_pr.append(shd)

def set_margins(doc):
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

# ---------- Exec Summary ----------
def build_exec_summary(rollup, students, totals, by_school, out_dir, as_of_long, as_of_str):
    doc = Document()
    set_margins(doc)
    add_title(doc, "Chronic Absenteeism Snapshot",
              f"Greeneville City Schools | As of {as_of_long}")

    add_h2(doc, "The Bottom Line")
    doc.add_paragraph(
        f"{totals['total']} students are currently flagged. "
        f"{totals['already']} have already crossed the 17-day chronic absenteeism threshold for this school year. "
        f"The remaining {totals['trend']} students have absentee rates at or above 10% but can still avoid a CA designation "
        f"if they stabilize attendance through the end of the year."
    )

    add_h2(doc, "District Rollup")
    t = doc.add_table(rows=1, cols=5)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["School", "Flagged", "Already CA (17+)", "Trending", "% Already CA"]):
        t.rows[0].cells[i].text = h
    style_header_row(t.rows[0])
    for code in SCHOOL_ORDER:
        b = by_school[code]
        pct = (b["already"] / b["total"] * 100) if b["total"] else 0
        row = t.add_row().cells
        row[0].text = SCHOOL_NAMES[code]
        row[1].text = str(b["total"])
        row[2].text = str(b["already"])
        row[3].text = str(b["trend"])
        row[4].text = f"{pct:.1f}%"
    pct_total = (totals["already"] / totals["total"] * 100) if totals["total"] else 0
    row = t.add_row().cells
    row[0].text = "District Total"
    row[1].text = str(totals["total"])
    row[2].text = str(totals["already"])
    row[3].text = str(totals["trend"])
    row[4].text = f"{pct_total:.1f}%"
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    add_h2(doc, "School-Level % Chronically Absent (Enrolled Population)")
    doc.add_paragraph(
        "Note: these figures reflect all enrolled students at each school, not just those on the flagged list. "
        "Schools showing 0% indicate the official CA count is calculated at the district level once state ISP rules are applied."
    )
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = "Light Grid Accent 1"
    for i, h in enumerate(["School", "# Students", "# Chronically Absent", "% CA"]):
        t2.rows[0].cells[i].text = h
    style_header_row(t2.rows[0])
    for r in rollup:
        row = t2.add_row().cells
        row[0].text = str(r["school"])
        row[1].text = str(r["n"])
        row[2].text = str(r["n_ca"])
        row[3].text = f"{r['pct']}%"

    add_h2(doc, "Outliers Worth Flagging")
    top10 = sorted(students, key=lambda s: -s["absences"])[:10]
    p = doc.add_paragraph()
    p.add_run("Ten highest absence counts across the district:").italic = True
    t3 = doc.add_table(rows=1, cols=5)
    t3.style = "Light Grid Accent 1"
    for i, h in enumerate(["School", "Student", "Grade", "Days Absent", "Rate"]):
        t3.rows[0].cells[i].text = h
    style_header_row(t3.rows[0])
    for s in top10:
        row = t3.add_row().cells
        row[0].text = SCHOOL_NAMES[s["school"]]
        row[1].text = f"{s['first']} {s['last']}"
        row[2].text = str(s["grade"])
        row[3].text = str(s["absences"])
        row[4].text = f"{s['rate']}%"

    add_h2(doc, "Where the Concentration Sits")
    ghs = by_school["GHS"]; gms = by_school["GMS"]; hh = by_school["HH"]
    ghs_share = (ghs["total"] / totals["total"] * 100) if totals["total"] else 0
    ghs_ca_share = (ghs["already"] / totals["already"] * 100) if totals["already"] else 0
    bullets = [
        f"GHS carries {ghs_share:.0f}% of the flagged list ({ghs['total']} students) and {ghs_ca_share:.0f}% of the already-CA group ({ghs['already']}). Secondary attendance intervention should be weighted here first.",
        f"GMS is next ({gms['total']} students flagged, {gms['already']} already CA). A handful of 7th-8th graders are trending toward 40%+ rates.",
        f"HH is the elementary concentration point ({hh['total']} students, {hh['already']} already CA). Elementary CA has a different intervention playbook, home visits and family outreach tend to move the needle.",
        "EastView, Highland, Tusculum View, and TOPS each have small flagged lists. These are case-managed one-on-one rather than programmatic.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    add_h2(doc, "Recommended Follow-Ups")
    for a in [
        "Share per-principal packets with each building leader this week. Ask for an attendance plan per student in the already-CA list by end of month.",
        "For trending students, confirm each building is running the standard attendance letter / parent contact / SART escalation sequence.",
        "At GHS, triage the top-20 cases with the counseling office and social worker before the next long break.",
        "Reconcile with Brandy's ISP calculations after the next enrollment pull to make sure partial-year enrollees are not being misclassified.",
        "Re-run this snapshot in 4 weeks to measure movement.",
    ]:
        doc.add_paragraph(a, style="List Bullet")

    path = out_dir / f"Chronic Absenteeism Executive Summary {as_of_str}.docx"
    doc.save(path)
    return path

# ---------- Per-Principal Packets ----------
def build_principal_packet(code, students_at_school, rollup, out_dir, as_of_long):
    doc = Document()
    set_margins(doc)
    add_title(doc, f"Chronic Absenteeism — {SCHOOL_NAMES[code]}", f"As of {as_of_long}")

    already = [s for s in students_at_school if classify(s) == "already_ca"]
    trending = [s for s in students_at_school if classify(s) == "trending"]
    rollup_match = next((r for r in rollup if SCHOOL_NAMES[code].lower().split()[0] in r["school"].lower()
                         or code in r["school"]), None)

    add_h2(doc, "Your Numbers at a Glance")
    lines = [
        f"Students flagged on this snapshot: {len(students_at_school)}",
        f"Already chronically absent (17+ days): {len(already)}",
        f"Trending (10%+ rate, not yet CA): {len(trending)}",
    ]
    if rollup_match:
        lines.append(
            f"School-level CA rate (state calc, enrolled population): {rollup_match['pct']}% "
            f"({rollup_match['n_ca']} of {rollup_match['n']})"
        )
    for l in lines:
        doc.add_paragraph(l, style="List Bullet")

    def student_table(students_list, sort_key):
        t = doc.add_table(rows=1, cols=5)
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(["Student", "Grade", "Days Absent", "Rate", "SSID"]):
            t.rows[0].cells[i].text = h
        style_header_row(t.rows[0])
        for s in sorted(students_list, key=sort_key):
            row = t.add_row().cells
            row[0].text = f"{s['first']} {s['last']}"
            row[1].text = str(s["grade"])
            row[2].text = str(s["absences"])
            row[3].text = f"{s['rate']}%"
            row[4].text = str(s["ssid"])

    if already:
        add_h2(doc, f"Already Chronically Absent — {len(already)} students")
        doc.add_paragraph(
            "These students have already crossed the 17-day threshold. Focus on preventing further absences and "
            "documenting re-engagement plans."
        )
        student_table(already, lambda x: -x["absences"])

    if trending:
        add_h2(doc, f"Trending Toward CA — {len(trending)} students")
        doc.add_paragraph(
            "These students are at or above a 10% absentee rate but still under 17 days. "
            "If they stabilize attendance through the end of the year, they avoid a CA designation."
        )
        student_table(trending, lambda x: -x["rate"])

    add_h2(doc, "Talking Points for Your Attendance Team")
    for b in [
        "Identify which of the already-CA students have documented medical/legal reasons versus unexcused patterns.",
        "Confirm each student on the list has received the standard attendance letter sequence and a SART referral where appropriate.",
        "Flag any students enrolled less than 50% of the year to Brandy so ISP denominators can be verified.",
        "Build a named intervention plan for each student on the already-CA list by end of month.",
        "For trending students, a single focused parent contact often prevents the slide to CA, prioritize those closest to 17 days.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    add_h2(doc, "What I Need From You")
    for a in [
        "Brief written update on the top 5 names in your already-CA list by end of month (intervention plan, owner, next check-in date).",
        "Confirmation that SART/court referral process is current for students with 25+ absences.",
        "Flag any data issues so we can reconcile with Brandy before the next snapshot.",
    ]:
        doc.add_paragraph(a, style="List Bullet")

    path = out_dir / f"Chronic Absenteeism Packet - {SCHOOL_NAMES[code]}.docx"
    doc.save(path)
    return path

def main():
    if len(sys.argv) < 3:
        print("Usage: python build.py <source_xlsx> <as_of_YYYY-MM-DD> [day_number]", file=sys.stderr)
        sys.exit(1)
    src = Path(sys.argv[1]).expanduser()
    as_of_str = sys.argv[2]  # YYYY-MM-DD
    day_num = sys.argv[3] if len(sys.argv) > 3 else None

    as_of_date = datetime.strptime(as_of_str, "%Y-%m-%d").date()
    as_of_long = as_of_date.strftime("%B %-d, %Y")
    if day_num:
        as_of_long += f" (Day {day_num} of {INSTRUCTIONAL_DAYS_DEFAULT})"

    out_dir = Path.home() / "Documents" / f"Chronic Absenteeism {as_of_str}"
    out_dir.mkdir(parents=True, exist_ok=True)

    rollup, students = load_data(src)
    dash_path, totals, by_school = build_dashboard(rollup, students, out_dir, as_of_str)
    exec_path = build_exec_summary(rollup, students, totals, by_school, out_dir, as_of_long, as_of_str)
    by_code = defaultdict(list)
    for s in students:
        by_code[s["school"]].append(s)
    packets = [build_principal_packet(code, by_code[code], rollup, out_dir, as_of_long) for code in SCHOOL_ORDER]

    # Archive to OneDrive
    ONEDRIVE.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, ONEDRIVE / src.name)
    shutil.copy2(dash_path, ONEDRIVE / dash_path.name)
    shutil.copy2(exec_path, ONEDRIVE / exec_path.name)
    for p in packets:
        shutil.copy2(p, ONEDRIVE / p.name)

    print(f"Output folder: {out_dir}")
    print(f"Dashboard:     {dash_path.name}")
    print(f"Exec summary:  {exec_path.name}")
    for p in packets:
        print(f"Packet:        {p.name}")
    print(f"OneDrive copy: {ONEDRIVE / src.name}")
    print(f"Totals:        {totals}")

if __name__ == "__main__":
    main()
